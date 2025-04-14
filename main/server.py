from flask import Flask, redirect, request, render_template, jsonify, flash, url_for, send_file
import sqlite3
import datetime
import os
import time
import threading
import docx
from collections import defaultdict
import playsound
import schedule

app = Flask(__name__)

USERNAME = "admin"
PASSWORD = "password123"

DB_NAME = 'threat_logs.db'
BLOCK_THRESHOLD = 10  # Number of requests from an IP before blocking
BLOCK_TIME = 60  # Seconds to block an IP

ATTACK_TYPES = {
    "high_traffic": "DDoS Attempt",
    "rapid_access": "Brute Force Attack"
}
app.secret_key = "1239123"
# Initialize database
def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS logs (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            ip TEXT,
                            page TEXT,
                            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS blocked_ips (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            ip TEXT UNIQUE,
                            block_time DATETIME DEFAULT CURRENT_TIMESTAMP)''')
        conn.commit()

def get_logs_from_db():
    conn = sqlite3.connect('threat_logs.db')  # Replace with your DB name
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM logs")  # Assuming `logs` table exists
    logs = [{"id": row[1], "message": row[2], "timestamp": row[3]} for row in cursor.fetchall()]
    conn.close()
    print(logs)
    return logs

# Function to fetch blocked IPs from the database
def get_blocked_ips_from_db():
    conn = sqlite3.connect('threat_logs.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM blocked_ips")  # Assuming `blocked_ips` table exists
    blocked_ips = [{"id": row[0], "ip_address": row[1], "reason": row[2]} for row in cursor.fetchall()]
    conn.close()
    return blocked_ips

# Store log
def store_log(ip, page):
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO logs (ip, page) VALUES (?, ?)", (ip, page))
        conn.commit()

# Check if IP is blocked
def is_ip_blocked(ip):
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM blocked_ips WHERE ip = ?", (ip,))
        return cursor.fetchone() is not None

# Block IP
def block_ip(ip):
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT OR IGNORE INTO blocked_ips (ip) VALUES (?)", (ip,))
        conn.commit()

    play_alert_sound()
    generate_report(ip)

# Play alert sound
def play_alert_sound():
    playsound.playsound("alert.mp3")

# Generate threat report for an IP
def generate_report(ip):
    doc = docx.Document()
    doc.add_heading(f"Threat Report for {ip}", level=1)
    
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT page, timestamp FROM logs WHERE ip = ?", (ip,))
        logs = cursor.fetchall()
        
        if logs:
            doc.add_paragraph(f"Total Requests: {len(logs)}")
            attack_type = "high_traffic" if len(logs) > 20 else "rapid_access"
            risk_level = "High" if len(logs) > 20 else "Intermediate"
            
            doc.add_paragraph(f"Attack Type: {ATTACK_TYPES[attack_type]}")
            doc.add_paragraph(f"Risk Level: {risk_level}")
            
            for log in logs:
                doc.add_paragraph(f"Page: {log[0]}, Timestamp: {log[1]}")

    if not os.path.exists("reports"):
        os.makedirs("reports")
    doc.save(f"reports/{ip}_report.docx")

# Generate daily server report
def generate_daily_report():
    doc = docx.Document()
    doc.add_heading("Daily Server Report", level=1)
    
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT ip, COUNT(*) FROM logs WHERE timestamp > datetime('now', '-1 day') GROUP BY ip")
        logs = cursor.fetchall()
        
        for log in logs:
            doc.add_paragraph(f"IP: {log[0]}, Total Requests: {log[1]}")

    if not os.path.exists("reports"):
        os.makedirs("reports")
    doc.save(f"reports/daily_report_{datetime.date.today()}.docx")

# Schedule daily report
def schedule_daily_report():
    schedule.every().day.at("23:59").do(generate_daily_report)
    while True:
        schedule.run_pending()
        time.sleep(60)

@app.route('/')
def home():
    return render_template('index.html')

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        if username == USERNAME and password == PASSWORD:
            return redirect(url_for("home"))
        else:
            flash("❌ Invalid username or password!", "danger")

    return render_template("login.html")

@app.route('/log', methods=['POST'])
def log_request():
    ip = request.remote_addr
    page = request.json.get("page", "/")
    
    if is_ip_blocked(ip):
        return jsonify({"message": "Access denied, IP blocked!"}), 403
    
    store_log(ip, page)
    
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM logs WHERE ip = ? AND timestamp > datetime('now', '-1 minute')", (ip,))
        count = cursor.fetchone()[0]
        
        if count > BLOCK_THRESHOLD:
            block_ip(ip)
            return jsonify({"message": "IP blocked due to high traffic!"}), 403
    
    return jsonify({"message": "Log stored successfully"})

@app.route('/unblock/<ip>', methods=['POST'])
def unblock_ip(ip):
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM blocked_ips WHERE ip = ?", (ip,))
        conn.commit()
    return jsonify({"message": f"IP {ip} unblocked successfully"}), 200



@app.route('/get_logs', methods=['GET'])
def get_logs():
    return jsonify({"logs": get_logs_from_db()})

@app.route('/get_blocked_ips', methods=['GET'])
def get_blocked_ips():
    return jsonify({"blocked_ips": get_blocked_ips_from_db()})

@app.route('/video.mp4')
def serve_video():
    return send_file("video.mp4", mimetype="video/mp4")
if __name__ == '__main__':
    init_db()
    threading.Thread(target=schedule_daily_report, daemon=True).start()
    app.run(debug=True)
